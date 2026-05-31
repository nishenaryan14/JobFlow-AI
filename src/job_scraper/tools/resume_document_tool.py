"""
Resume Document Generator Tool
================================
A tool that takes a formatted Markdown resume string and renders it
into a professional PDF using ReportLab, saving the file to /tmp/.

Called directly by the resume_enhancement graph node — no framework dependency.
"""

import re
import os
import tempfile
from pathlib import Path
from pydantic import BaseModel, Field


class ResumeDocumentInput(BaseModel):
    """Input schema for the Resume Document Generator tool."""
    markdown_resume: str = Field(
        description="The complete enhanced resume text in Markdown format. "
                    "Must include Name as H1, sections as H2, bullet points for experience.")
    output_format: str = Field(
        default="pdf",
        description="Output format: 'pdf' or 'docx'. Defaults to pdf.")


class ResumeDocumentTool(BaseModel):
    """
    Generates a professionally formatted PDF or DOCX resume from Markdown text.
    Use this tool AFTER you have written the complete enhanced resume content.
    """
    name: str = "resume_document_generator"
    description: str = (
        "Converts a Markdown-formatted resume into a professional PDF or DOCX file. "
        "Call this once you have the complete enhanced resume ready. "
        "Returns the file path where the document was saved."
    )
    args_schema: type[BaseModel] = ResumeDocumentInput

    def _run(self, markdown_resume: str, output_format: str = "pdf") -> str:
        try:
            out_dir = Path(tempfile.gettempdir()) / "jobflow_resumes"
            out_dir.mkdir(parents=True, exist_ok=True)

            if output_format.lower() == "docx":
                return self._generate_docx(markdown_resume, out_dir)
            else:
                return self._generate_pdf(markdown_resume, out_dir)
        except Exception as e:
            return f"ERROR: Document generation failed — {e}"

    # ─── PDF Generation via ReportLab ────────────────────────────────────────

    def _generate_pdf(self, markdown: str, out_dir: Path) -> str:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
            ListFlowable, ListItem
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT

        out_path = out_dir / "enhanced_resume.pdf"

        doc = SimpleDocTemplate(
            str(out_path),
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=18 * mm,
            bottomMargin=18 * mm,
        )

        # ── Style sheet ──
        styles = getSampleStyleSheet()

        name_style = ParagraphStyle(
            "Name",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=2,
            textColor=colors.HexColor("#111111"),
        )
        contact_style = ParagraphStyle(
            "Contact",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9,
            alignment=TA_CENTER,
            spaceAfter=6,
            textColor=colors.HexColor("#555555"),
        )
        section_style = ParagraphStyle(
            "Section",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            spaceBefore=10,
            spaceAfter=3,
            textColor=colors.HexColor("#111111"),
            borderPadding=(0, 0, 2, 0),
        )
        job_title_style = ParagraphStyle(
            "JobTitle",
            parent=styles["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            spaceBefore=5,
            spaceAfter=1,
            textColor=colors.HexColor("#1a1a1a"),
        )
        job_meta_style = ParagraphStyle(
            "JobMeta",
            parent=styles["Normal"],
            fontName="Helvetica-Oblique",
            fontSize=9,
            spaceAfter=2,
            textColor=colors.HexColor("#555555"),
        )
        body_style = ParagraphStyle(
            "Body",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            spaceAfter=3,
            leading=13,
            textColor=colors.HexColor("#2a2a2a"),
        )
        bullet_style = ParagraphStyle(
            "Bullet",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=9.5,
            spaceAfter=2,
            leading=13,
            leftIndent=12,
            textColor=colors.HexColor("#2a2a2a"),
        )

        story = []
        lines = markdown.strip().split("\n")
        i = 0

        # Helper: clean inline markdown from a line
        def clean(text: str) -> str:
            text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            text = re.sub(r"\*\*\*(.+?)\*\*\*", r"<b><i>\1</i></b>", text)
            text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
            text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
            text = re.sub(r"`(.+?)`", r"<font name='Courier'>\1</font>", text)
            return text

        in_bullet_group = False
        bullet_items = []

        def flush_bullets():
            nonlocal in_bullet_group, bullet_items
            if bullet_items:
                for b in bullet_items:
                    story.append(Paragraph(f"• {b}", bullet_style))
            bullet_items = []
            in_bullet_group = False

        while i < len(lines):
            line = lines[i]
            stripped = line.strip()

            # Skip blank lines (flush bullet groups)
            if not stripped:
                flush_bullets()
                story.append(Spacer(1, 2))
                i += 1
                continue

            # H1 → Name
            if stripped.startswith("# "):
                flush_bullets()
                text = stripped[2:].strip()
                story.append(Paragraph(clean(text), name_style))
                story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#111111"), spaceAfter=3))
                i += 1
                continue

            # H2 → Section heading
            if stripped.startswith("## "):
                flush_bullets()
                text = stripped[3:].strip().upper()
                story.append(Spacer(1, 4))
                story.append(Paragraph(clean(text), section_style))
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#bbbbbb"), spaceAfter=3))
                i += 1
                continue

            # H3 → Job/project title
            if stripped.startswith("### "):
                flush_bullets()
                text = stripped[4:].strip()
                story.append(Paragraph(clean(text), job_title_style))
                i += 1
                continue

            # Horizontal rule
            if re.match(r"^-{3,}$", stripped):
                flush_bullets()
                story.append(HRFlowable(width="100%", thickness=0.5, color=colors.HexColor("#dddddd"), spaceAfter=4))
                i += 1
                continue

            # Bullet point
            if stripped.startswith(("- ", "* ", "• ")):
                text = stripped[2:].strip() if stripped[1] == " " else stripped[1:].strip()
                bullet_items.append(clean(text))
                i += 1
                continue

            # Regular text / contact info / italic meta
            flush_bullets()
            # Check if this looks like a contact line (has @, phone, linkedin)
            if re.search(r"[@|📧📱🔗]|linkedin|github|\+\d", stripped):
                story.append(Paragraph(clean(stripped), contact_style))
            elif stripped.startswith("_") and stripped.endswith("_"):
                story.append(Paragraph(clean(stripped.strip("_")), job_meta_style))
            else:
                story.append(Paragraph(clean(stripped), body_style))
            i += 1

        flush_bullets()

        doc.build(story)
        return str(out_path)

    # ─── DOCX Generation via python-docx ─────────────────────────────────────

    def _generate_docx(self, markdown: str, out_dir: Path) -> str:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        out_path = out_dir / "enhanced_resume.docx"
        doc = Document()

        # Page margins
        for section in doc.sections:
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)
            section.top_margin = Inches(0.7)
            section.bottom_margin = Inches(0.7)

        def add_hr(doc):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)

        def clean_inline(text: str) -> list:
            """Returns list of (run_text, bold, italic) tuples."""
            parts = []
            pattern = re.compile(r"\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+)")
            for m in pattern.finditer(text):
                if m.group(1):
                    parts.append((m.group(1), True, True))
                elif m.group(2):
                    parts.append((m.group(2), True, False))
                elif m.group(3):
                    parts.append((m.group(3), False, True))
                elif m.group(4):
                    parts.append((m.group(4), False, False))
            return parts

        def add_para(text: str, bold=False, italic=False, size=10, align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=4):
            p = doc.add_paragraph()
            p.alignment = align
            p.paragraph_format.space_before = Pt(space_before)
            p.paragraph_format.space_after = Pt(space_after)
            for chunk, b, it in clean_inline(text):
                run = p.add_run(chunk)
                run.bold = bold or b
                run.italic = italic or it
                run.font.size = Pt(size)
                if color:
                    run.font.color.rgb = RGBColor(*color)
            return p

        lines = markdown.strip().split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                continue

            if stripped.startswith("# "):
                p = add_para(stripped[2:].strip(), bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
                add_hr(doc)
            elif stripped.startswith("## "):
                add_para(stripped[3:].strip().upper(), bold=True, size=10, color=(17, 17, 17), space_before=8, space_after=2)
                add_hr(doc)
            elif stripped.startswith("### "):
                add_para(stripped[4:].strip(), bold=True, size=10, space_before=5, space_after=1)
            elif re.match(r"^-{3,}$", stripped):
                add_hr(doc)
            elif stripped.startswith(("- ", "* ", "• ")):
                text = stripped[2:] if stripped[1] == " " else stripped[1:]
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(2)
                for chunk, b, it in clean_inline(text.strip()):
                    run = p.add_run(chunk)
                    run.bold = b
                    run.italic = it
                    run.font.size = Pt(9.5)
            elif re.search(r"@|linkedin|github|\+\d", stripped):
                add_para(stripped, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=(85, 85, 85), space_after=2)
            else:
                add_para(stripped, size=9.5, space_after=3)

        doc.save(str(out_path))
        return str(out_path)
