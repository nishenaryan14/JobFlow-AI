"""Resume Template Renderer — Converts structured resume data into HTML/PDF/DOCX.

Templates:
  • Modern  — Clean sans-serif (Inter), blue accent, two-column skills
  • Classic — Traditional serif (Georgia), single column, bold rules
  • Minimal — Ultra-clean, lots of whitespace, monochrome
  • ATS     — Zero graphics, standard headers, Times New Roman

All HTML uses inline CSS for reliable PDF conversion.
"""

import io
from html import escape
from typing import Any, Dict, List


# ── Helper utilities ──────────────────────────────────────────────────────────

def _e(text: Any) -> str:
    """HTML-escape a value, handling non-strings."""
    return escape(str(text)) if text else ""


def _has(data: dict, key: str) -> bool:
    """Check if a key exists and has non-empty content."""
    val = data.get(key)
    if val is None:
        return False
    if isinstance(val, str):
        return bool(val.strip())
    if isinstance(val, list):
        return len(val) > 0
    if isinstance(val, dict):
        return any(v for v in val.values() if v)
    return bool(val)


def _contact(data: dict) -> dict:
    """Extract contact info, handling nested or flat structure."""
    contact = data.get("contact", {})
    if isinstance(contact, dict):
        return contact
    return data


# ── Modern Template ───────────────────────────────────────────────────────────

def render_modern_html(data: dict) -> str:
    """Clean sans-serif (Inter), subtle blue accent (#3772ff), two-column skills."""
    contact = _contact(data)
    html_parts = [f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{_e(contact.get('name', 'Resume'))}</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
@page {{ size: A4; margin: 15mm 18mm; }}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: 'Inter', sans-serif; font-size: 10pt; line-height: 1.5; color: #1a1a2e; }}
.container {{ max-width: 210mm; margin: 0 auto; padding: 20px 0; }}
h1 {{ font-size: 22pt; font-weight: 700; color: #3772ff; margin-bottom: 4px; }}
h2 {{ font-size: 11pt; font-weight: 600; color: #3772ff; text-transform: uppercase;
      letter-spacing: 1.5px; border-bottom: 2px solid #3772ff; padding-bottom: 4px;
      margin: 18px 0 10px 0; }}
.contact-line {{ font-size: 9pt; color: #555; margin-bottom: 2px; }}
.contact-line a {{ color: #3772ff; text-decoration: none; }}
.summary {{ font-size: 10pt; color: #333; margin: 10px 0 16px 0; line-height: 1.6; }}
.exp-header {{ display: flex; justify-content: space-between; align-items: baseline; margin-bottom: 2px; }}
.exp-title {{ font-weight: 600; font-size: 10.5pt; }}
.exp-company {{ font-weight: 500; color: #555; }}
.exp-date {{ font-size: 9pt; color: #777; white-space: nowrap; }}
.exp-location {{ font-size: 9pt; color: #888; }}
ul {{ padding-left: 18px; margin: 4px 0 12px 0; }}
li {{ margin-bottom: 3px; font-size: 9.5pt; }}
.skills-grid {{ display: grid; grid-template-columns: 1fr 1fr; gap: 6px 20px; }}
.skill-cat {{ font-weight: 600; font-size: 9.5pt; color: #3772ff; }}
.skill-list {{ font-size: 9.5pt; color: #333; }}
.edu-header {{ display: flex; justify-content: space-between; align-items: baseline; }}
.edu-degree {{ font-weight: 600; font-size: 10pt; }}
.edu-year {{ font-size: 9pt; color: #777; }}
.edu-institution {{ font-size: 9.5pt; color: #555; }}
.project-name {{ font-weight: 600; font-size: 10pt; }}
.project-tech {{ font-size: 8.5pt; color: #3772ff; font-style: italic; }}
.cert-list, .lang-list {{ font-size: 9.5pt; columns: 2; column-gap: 20px; }}
</style>
</head>
<body>
<div class="container">
"""]

    # Header
    if contact.get("name"):
        html_parts.append(f'<h1>{_e(contact["name"])}</h1>')
    contact_items = []
    for key in ["email", "phone", "linkedin", "github", "portfolio"]:
        val = contact.get(key, "")
        if val:
            if key in ("linkedin", "github", "portfolio") and val.startswith("http"):
                contact_items.append(f'<a href="{_e(val)}">{_e(val)}</a>')
            else:
                contact_items.append(_e(val))
    if contact_items:
        html_parts.append(f'<div class="contact-line">{" &nbsp;|&nbsp; ".join(contact_items)}</div>')

    # Summary
    if _has(data, "summary"):
        html_parts.append(f'<div class="summary">{_e(data["summary"])}</div>')

    # Experience
    experience = data.get("experience", [])
    if experience:
        html_parts.append('<h2>Experience</h2>')
        for exp in experience:
            html_parts.append('<div style="margin-bottom: 12px;">')
            html_parts.append('<div class="exp-header">')
            html_parts.append(f'<span class="exp-title">{_e(exp.get("title", ""))}</span>')
            dates = f'{_e(exp.get("start_date", ""))} – {_e(exp.get("end_date", ""))}'
            html_parts.append(f'<span class="exp-date">{dates}</span>')
            html_parts.append('</div>')
            company_loc = _e(exp.get("company", ""))
            if exp.get("location"):
                company_loc += f' &nbsp;·&nbsp; {_e(exp["location"])}'
            html_parts.append(f'<div class="exp-company">{company_loc}</div>')
            bullets = exp.get("bullets", [])
            if bullets:
                html_parts.append('<ul>')
                for b in bullets:
                    html_parts.append(f'<li>{_e(b)}</li>')
                html_parts.append('</ul>')
            html_parts.append('</div>')

    # Education
    education = data.get("education", [])
    if education:
        html_parts.append('<h2>Education</h2>')
        for edu in education:
            html_parts.append('<div style="margin-bottom: 8px;">')
            html_parts.append('<div class="edu-header">')
            html_parts.append(f'<span class="edu-degree">{_e(edu.get("degree", ""))}</span>')
            html_parts.append(f'<span class="edu-year">{_e(edu.get("year", ""))}</span>')
            html_parts.append('</div>')
            inst_loc = _e(edu.get("institution", ""))
            if edu.get("location"):
                inst_loc += f' &nbsp;·&nbsp; {_e(edu["location"])}'
            html_parts.append(f'<div class="edu-institution">{inst_loc}</div>')
            extras = []
            if edu.get("gpa"):
                extras.append(f'GPA: {_e(edu["gpa"])}')
            if edu.get("honors"):
                extras.append(_e(edu["honors"]))
            if extras:
                html_parts.append(f'<div style="font-size:9pt;color:#777;">{" &nbsp;|&nbsp; ".join(extras)}</div>')
            html_parts.append('</div>')

    # Skills
    skills = data.get("skills", [])
    if skills:
        html_parts.append('<h2>Skills</h2>')
        html_parts.append('<div class="skills-grid">')
        for cat in skills:
            cat_name = cat.get("name", "")
            cat_skills = cat.get("skills", [])
            if cat_skills:
                html_parts.append(f'<div><span class="skill-cat">{_e(cat_name)}:</span> '
                                  f'<span class="skill-list">{", ".join(_e(s) for s in cat_skills)}</span></div>')
        html_parts.append('</div>')

    # Projects
    projects = data.get("projects", [])
    if projects:
        html_parts.append('<h2>Projects</h2>')
        for proj in projects:
            html_parts.append('<div style="margin-bottom: 10px;">')
            name_line = f'<span class="project-name">{_e(proj.get("name", ""))}</span>'
            techs = proj.get("technologies", [])
            if techs:
                name_line += f' <span class="project-tech">({", ".join(_e(t) for t in techs)})</span>'
            html_parts.append(name_line)
            if proj.get("description"):
                html_parts.append(f'<div style="font-size:9.5pt;color:#444;">{_e(proj["description"])}</div>')
            bullets = proj.get("bullets", [])
            if bullets:
                html_parts.append('<ul>')
                for b in bullets:
                    html_parts.append(f'<li>{_e(b)}</li>')
                html_parts.append('</ul>')
            html_parts.append('</div>')

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        html_parts.append('<h2>Certifications</h2>')
        html_parts.append('<div class="cert-list">')
        for c in certs:
            html_parts.append(f'<div>• {_e(c)}</div>')
        html_parts.append('</div>')

    # Languages
    langs = data.get("languages", [])
    if langs:
        html_parts.append('<h2>Languages</h2>')
        html_parts.append(f'<div class="lang-list">{" &nbsp;|&nbsp; ".join(_e(l) for l in langs)}</div>')

    html_parts.append('</div></body></html>')
    return "\n".join(html_parts)


# ── Classic Template ──────────────────────────────────────────────────────────

def render_classic_html(data: dict) -> str:
    """Traditional serif (Georgia), single column, bold horizontal rules."""
    contact = _contact(data)
    html_parts = [f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{_e(contact.get('name', 'Resume'))}</title>
<style>
@page {{ size: A4; margin: 20mm; }}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: Georgia, 'Times New Roman', serif; font-size: 10.5pt; line-height: 1.5; color: #222; }}
.container {{ max-width: 210mm; margin: 0 auto; padding: 20px 0; }}
h1 {{ font-size: 20pt; font-weight: 700; text-align: center; margin-bottom: 4px; color: #111; }}
h2 {{ font-size: 12pt; font-weight: 700; text-transform: uppercase; color: #111;
      border-bottom: 2.5px solid #111; padding-bottom: 3px; margin: 20px 0 10px 0; }}
.contact-center {{ text-align: center; font-size: 9.5pt; color: #444; margin-bottom: 14px; }}
.contact-center a {{ color: #444; text-decoration: none; }}
.summary {{ font-size: 10.5pt; text-align: justify; margin-bottom: 14px; color: #333; font-style: italic; }}
.entry {{ margin-bottom: 14px; }}
.entry-top {{ display: flex; justify-content: space-between; }}
.entry-title {{ font-weight: 700; font-size: 10.5pt; }}
.entry-date {{ font-size: 9.5pt; color: #555; }}
.entry-sub {{ font-style: italic; color: #555; font-size: 10pt; }}
ul {{ padding-left: 20px; margin: 4px 0 0 0; }}
li {{ margin-bottom: 3px; }}
.skills-block {{ margin-bottom: 6px; }}
.skills-label {{ font-weight: 700; }}
</style>
</head>
<body>
<div class="container">
"""]

    # Header
    if contact.get("name"):
        html_parts.append(f'<h1>{_e(contact["name"])}</h1>')
    contact_items = []
    for key in ["email", "phone", "linkedin", "github", "portfolio"]:
        val = contact.get(key, "")
        if val:
            contact_items.append(_e(val))
    if contact_items:
        html_parts.append(f'<div class="contact-center">{" &nbsp;·&nbsp; ".join(contact_items)}</div>')

    # Summary
    if _has(data, "summary"):
        html_parts.append(f'<div class="summary">{_e(data["summary"])}</div>')

    # Experience
    experience = data.get("experience", [])
    if experience:
        html_parts.append('<h2>Professional Experience</h2>')
        for exp in experience:
            html_parts.append('<div class="entry">')
            html_parts.append('<div class="entry-top">')
            html_parts.append(f'<span class="entry-title">{_e(exp.get("title", ""))}</span>')
            dates = f'{_e(exp.get("start_date", ""))} – {_e(exp.get("end_date", ""))}'
            html_parts.append(f'<span class="entry-date">{dates}</span>')
            html_parts.append('</div>')
            sub = _e(exp.get("company", ""))
            if exp.get("location"):
                sub += f', {_e(exp["location"])}'
            html_parts.append(f'<div class="entry-sub">{sub}</div>')
            bullets = exp.get("bullets", [])
            if bullets:
                html_parts.append('<ul>')
                for b in bullets:
                    html_parts.append(f'<li>{_e(b)}</li>')
                html_parts.append('</ul>')
            html_parts.append('</div>')

    # Education
    education = data.get("education", [])
    if education:
        html_parts.append('<h2>Education</h2>')
        for edu in education:
            html_parts.append('<div class="entry">')
            html_parts.append('<div class="entry-top">')
            html_parts.append(f'<span class="entry-title">{_e(edu.get("degree", ""))}</span>')
            html_parts.append(f'<span class="entry-date">{_e(edu.get("year", ""))}</span>')
            html_parts.append('</div>')
            sub = _e(edu.get("institution", ""))
            if edu.get("location"):
                sub += f', {_e(edu["location"])}'
            html_parts.append(f'<div class="entry-sub">{sub}</div>')
            extras = []
            if edu.get("gpa"):
                extras.append(f'GPA: {_e(edu["gpa"])}')
            if edu.get("honors"):
                extras.append(_e(edu["honors"]))
            if extras:
                html_parts.append(f'<div style="font-size:9.5pt;color:#555;">{" | ".join(extras)}</div>')
            html_parts.append('</div>')

    # Skills
    skills = data.get("skills", [])
    if skills:
        html_parts.append('<h2>Skills</h2>')
        for cat in skills:
            cat_name = cat.get("name", "")
            cat_skills = cat.get("skills", [])
            if cat_skills:
                html_parts.append(f'<div class="skills-block"><span class="skills-label">'
                                  f'{_e(cat_name)}:</span> {", ".join(_e(s) for s in cat_skills)}</div>')

    # Projects
    projects = data.get("projects", [])
    if projects:
        html_parts.append('<h2>Projects</h2>')
        for proj in projects:
            html_parts.append('<div class="entry">')
            name = _e(proj.get("name", ""))
            techs = proj.get("technologies", [])
            tech_str = f' ({", ".join(_e(t) for t in techs)})' if techs else ""
            html_parts.append(f'<div class="entry-title">{name}{tech_str}</div>')
            if proj.get("description"):
                html_parts.append(f'<div class="entry-sub">{_e(proj["description"])}</div>')
            bullets = proj.get("bullets", [])
            if bullets:
                html_parts.append('<ul>')
                for b in bullets:
                    html_parts.append(f'<li>{_e(b)}</li>')
                html_parts.append('</ul>')
            html_parts.append('</div>')

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        html_parts.append('<h2>Certifications</h2>')
        html_parts.append('<ul>')
        for c in certs:
            html_parts.append(f'<li>{_e(c)}</li>')
        html_parts.append('</ul>')

    # Languages
    langs = data.get("languages", [])
    if langs:
        html_parts.append('<h2>Languages</h2>')
        html_parts.append(f'<div>{", ".join(_e(l) for l in langs)}</div>')

    html_parts.append('</div></body></html>')
    return "\n".join(html_parts)


# ── Minimal Template ─────────────────────────────────────────────────────────

def render_minimal_html(data: dict) -> str:
    """Ultra-clean, lots of whitespace, thin borders, monochrome with #333 accent."""
    contact = _contact(data)
    html_parts = [f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{_e(contact.get('name', 'Resume'))}</title>
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600&display=swap');
@page {{ size: A4; margin: 18mm 22mm; }}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: 'Inter', Helvetica, Arial, sans-serif; font-size: 9.5pt; line-height: 1.6;
        color: #333; }}
.container {{ max-width: 210mm; margin: 0 auto; padding: 24px 0; }}
h1 {{ font-size: 18pt; font-weight: 300; letter-spacing: 3px; text-transform: uppercase;
      color: #333; margin-bottom: 6px; }}
h2 {{ font-size: 9pt; font-weight: 500; letter-spacing: 2.5px; text-transform: uppercase;
      color: #333; border-bottom: 0.5px solid #ccc; padding-bottom: 4px;
      margin: 22px 0 12px 0; }}
.contact-line {{ font-size: 8.5pt; color: #888; letter-spacing: 0.5px; margin-bottom: 16px; }}
.contact-line a {{ color: #555; text-decoration: none; }}
.summary {{ font-size: 9.5pt; color: #555; margin-bottom: 18px; line-height: 1.7; }}
.item {{ margin-bottom: 14px; }}
.item-top {{ display: flex; justify-content: space-between; align-items: baseline; }}
.item-title {{ font-weight: 500; font-size: 10pt; color: #222; }}
.item-date {{ font-size: 8.5pt; color: #999; }}
.item-sub {{ font-size: 9pt; color: #777; margin-top: 1px; }}
ul {{ padding-left: 16px; margin: 4px 0 0 0; }}
li {{ margin-bottom: 2px; font-size: 9pt; color: #444; }}
.skills-row {{ font-size: 9pt; margin-bottom: 4px; }}
.skills-label {{ font-weight: 500; color: #333; }}
.tags {{ font-size: 8.5pt; color: #666; }}
</style>
</head>
<body>
<div class="container">
"""]

    # Header
    if contact.get("name"):
        html_parts.append(f'<h1>{_e(contact["name"])}</h1>')
    contact_items = []
    for key in ["email", "phone", "linkedin", "github", "portfolio"]:
        val = contact.get(key, "")
        if val:
            if key in ("linkedin", "github", "portfolio") and val.startswith("http"):
                contact_items.append(f'<a href="{_e(val)}">{_e(val)}</a>')
            else:
                contact_items.append(_e(val))
    if contact_items:
        html_parts.append(f'<div class="contact-line">{" &nbsp;&nbsp;/&nbsp;&nbsp; ".join(contact_items)}</div>')

    # Summary
    if _has(data, "summary"):
        html_parts.append(f'<div class="summary">{_e(data["summary"])}</div>')

    # Experience
    experience = data.get("experience", [])
    if experience:
        html_parts.append('<h2>Experience</h2>')
        for exp in experience:
            html_parts.append('<div class="item">')
            html_parts.append('<div class="item-top">')
            html_parts.append(f'<span class="item-title">{_e(exp.get("title", ""))}</span>')
            dates = f'{_e(exp.get("start_date", ""))} – {_e(exp.get("end_date", ""))}'
            html_parts.append(f'<span class="item-date">{dates}</span>')
            html_parts.append('</div>')
            sub = _e(exp.get("company", ""))
            if exp.get("location"):
                sub += f' · {_e(exp["location"])}'
            html_parts.append(f'<div class="item-sub">{sub}</div>')
            bullets = exp.get("bullets", [])
            if bullets:
                html_parts.append('<ul>')
                for b in bullets:
                    html_parts.append(f'<li>{_e(b)}</li>')
                html_parts.append('</ul>')
            html_parts.append('</div>')

    # Education
    education = data.get("education", [])
    if education:
        html_parts.append('<h2>Education</h2>')
        for edu in education:
            html_parts.append('<div class="item">')
            html_parts.append('<div class="item-top">')
            html_parts.append(f'<span class="item-title">{_e(edu.get("degree", ""))}</span>')
            html_parts.append(f'<span class="item-date">{_e(edu.get("year", ""))}</span>')
            html_parts.append('</div>')
            sub = _e(edu.get("institution", ""))
            if edu.get("location"):
                sub += f' · {_e(edu["location"])}'
            html_parts.append(f'<div class="item-sub">{sub}</div>')
            html_parts.append('</div>')

    # Skills
    skills = data.get("skills", [])
    if skills:
        html_parts.append('<h2>Skills</h2>')
        for cat in skills:
            cat_name = cat.get("name", "")
            cat_skills = cat.get("skills", [])
            if cat_skills:
                html_parts.append(f'<div class="skills-row"><span class="skills-label">'
                                  f'{_e(cat_name)}</span> &nbsp;— &nbsp;'
                                  f'{", ".join(_e(s) for s in cat_skills)}</div>')

    # Projects
    projects = data.get("projects", [])
    if projects:
        html_parts.append('<h2>Projects</h2>')
        for proj in projects:
            html_parts.append('<div class="item">')
            html_parts.append(f'<div class="item-title">{_e(proj.get("name", ""))}</div>')
            techs = proj.get("technologies", [])
            if techs:
                html_parts.append(f'<div class="tags">{", ".join(_e(t) for t in techs)}</div>')
            if proj.get("description"):
                html_parts.append(f'<div class="item-sub">{_e(proj["description"])}</div>')
            bullets = proj.get("bullets", [])
            if bullets:
                html_parts.append('<ul>')
                for b in bullets:
                    html_parts.append(f'<li>{_e(b)}</li>')
                html_parts.append('</ul>')
            html_parts.append('</div>')

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        html_parts.append('<h2>Certifications</h2>')
        for c in certs:
            html_parts.append(f'<div style="font-size:9pt;margin-bottom:3px;">· {_e(c)}</div>')

    # Languages
    langs = data.get("languages", [])
    if langs:
        html_parts.append('<h2>Languages</h2>')
        html_parts.append(f'<div style="font-size:9pt;">{" &nbsp;/&nbsp; ".join(_e(l) for l in langs)}</div>')

    html_parts.append('</div></body></html>')
    return "\n".join(html_parts)


# ── ATS Template ──────────────────────────────────────────────────────────────

def render_ats_html(data: dict) -> str:
    """Zero graphics, standard section headers (all caps), Times New Roman."""
    contact = _contact(data)
    html_parts = [f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{_e(contact.get('name', 'Resume'))}</title>
<style>
@page {{ size: A4; margin: 18mm 20mm; }}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{ font-family: 'Times New Roman', Times, serif; font-size: 11pt; line-height: 1.4; color: #000; }}
.container {{ max-width: 210mm; margin: 0 auto; padding: 16px 0; }}
h1 {{ font-size: 16pt; font-weight: 700; text-align: center; margin-bottom: 4px; }}
h2 {{ font-size: 11pt; font-weight: 700; text-transform: uppercase;
      border-bottom: 1px solid #000; padding-bottom: 2px; margin: 16px 0 8px 0; }}
.contact-center {{ text-align: center; font-size: 10pt; margin-bottom: 12px; }}
.summary {{ font-size: 10.5pt; margin-bottom: 12px; }}
.entry {{ margin-bottom: 10px; }}
.entry-top {{ display: flex; justify-content: space-between; }}
.entry-title {{ font-weight: 700; }}
.entry-date {{ font-size: 10pt; }}
.entry-sub {{ font-size: 10.5pt; }}
ul {{ padding-left: 20px; margin: 3px 0 0 0; }}
li {{ margin-bottom: 2px; }}
.skills-line {{ margin-bottom: 4px; }}
.skills-label {{ font-weight: 700; }}
</style>
</head>
<body>
<div class="container">
"""]

    # Header
    if contact.get("name"):
        html_parts.append(f'<h1>{_e(contact["name"])}</h1>')
    contact_items = []
    for key in ["email", "phone", "linkedin", "github", "portfolio"]:
        val = contact.get(key, "")
        if val:
            contact_items.append(_e(val))
    if contact_items:
        html_parts.append(f'<div class="contact-center">{" | ".join(contact_items)}</div>')

    # Summary / Objective
    if _has(data, "summary"):
        html_parts.append('<h2>PROFESSIONAL SUMMARY</h2>')
        html_parts.append(f'<div class="summary">{_e(data["summary"])}</div>')

    # Experience
    experience = data.get("experience", [])
    if experience:
        html_parts.append('<h2>WORK EXPERIENCE</h2>')
        for exp in experience:
            html_parts.append('<div class="entry">')
            html_parts.append('<div class="entry-top">')
            title_company = _e(exp.get("title", ""))
            if exp.get("company"):
                title_company += f', {_e(exp["company"])}'
            if exp.get("location"):
                title_company += f', {_e(exp["location"])}'
            html_parts.append(f'<span class="entry-title">{title_company}</span>')
            dates = f'{_e(exp.get("start_date", ""))} – {_e(exp.get("end_date", ""))}'
            html_parts.append(f'<span class="entry-date">{dates}</span>')
            html_parts.append('</div>')
            bullets = exp.get("bullets", [])
            if bullets:
                html_parts.append('<ul>')
                for b in bullets:
                    html_parts.append(f'<li>{_e(b)}</li>')
                html_parts.append('</ul>')
            html_parts.append('</div>')

    # Education
    education = data.get("education", [])
    if education:
        html_parts.append('<h2>EDUCATION</h2>')
        for edu in education:
            html_parts.append('<div class="entry">')
            html_parts.append('<div class="entry-top">')
            degree_inst = _e(edu.get("degree", ""))
            if edu.get("institution"):
                degree_inst += f', {_e(edu["institution"])}'
            if edu.get("location"):
                degree_inst += f', {_e(edu["location"])}'
            html_parts.append(f'<span class="entry-title">{degree_inst}</span>')
            html_parts.append(f'<span class="entry-date">{_e(edu.get("year", ""))}</span>')
            html_parts.append('</div>')
            extras = []
            if edu.get("gpa"):
                extras.append(f'GPA: {_e(edu["gpa"])}')
            if edu.get("honors"):
                extras.append(_e(edu["honors"]))
            if extras:
                html_parts.append(f'<div class="entry-sub">{" | ".join(extras)}</div>')
            html_parts.append('</div>')

    # Skills — maximum keyword density
    skills = data.get("skills", [])
    if skills:
        html_parts.append('<h2>SKILLS</h2>')
        for cat in skills:
            cat_name = cat.get("name", "")
            cat_skills = cat.get("skills", [])
            if cat_skills:
                html_parts.append(f'<div class="skills-line"><span class="skills-label">'
                                  f'{_e(cat_name)}:</span> {", ".join(_e(s) for s in cat_skills)}</div>')

    # Projects
    projects = data.get("projects", [])
    if projects:
        html_parts.append('<h2>PROJECTS</h2>')
        for proj in projects:
            html_parts.append('<div class="entry">')
            name = _e(proj.get("name", ""))
            techs = proj.get("technologies", [])
            tech_str = f' ({", ".join(_e(t) for t in techs)})' if techs else ""
            html_parts.append(f'<div class="entry-title">{name}{tech_str}</div>')
            if proj.get("description"):
                html_parts.append(f'<div class="entry-sub">{_e(proj["description"])}</div>')
            bullets = proj.get("bullets", [])
            if bullets:
                html_parts.append('<ul>')
                for b in bullets:
                    html_parts.append(f'<li>{_e(b)}</li>')
                html_parts.append('</ul>')
            html_parts.append('</div>')

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        html_parts.append('<h2>CERTIFICATIONS</h2>')
        html_parts.append('<ul>')
        for c in certs:
            html_parts.append(f'<li>{_e(c)}</li>')
        html_parts.append('</ul>')

    # Languages
    langs = data.get("languages", [])
    if langs:
        html_parts.append('<h2>LANGUAGES</h2>')
        html_parts.append(f'<div>{", ".join(_e(l) for l in langs)}</div>')

    html_parts.append('</div></body></html>')
    return "\n".join(html_parts)


# ── Dispatcher ────────────────────────────────────────────────────────────────

def render_resume_html(data: dict, template: str = "modern") -> str:
    """Render structured resume data to HTML using the specified template."""
    renderers = {
        "modern": render_modern_html,
        "classic": render_classic_html,
        "minimal": render_minimal_html,
        "ats": render_ats_html,
    }
    return renderers.get(template, render_modern_html)(data)


# ── PDF Generator ─────────────────────────────────────────────────────────────

def generate_pdf(data: dict, template: str = "modern") -> bytes:
    """Convert structured resume data to PDF bytes via xhtml2pdf.

    xhtml2pdf is used over weasyprint because it has fewer system-level
    dependencies (no GTK/Cairo) and works reliably on Windows.
    """
    from xhtml2pdf import pisa

    html = render_resume_html(data, template)
    buffer = io.BytesIO()
    pisa_status = pisa.CreatePDF(io.StringIO(html), dest=buffer)

    if pisa_status.err:
        raise RuntimeError(f"PDF generation failed with {pisa_status.err} errors")

    return buffer.getvalue()


# ── DOCX Generator ────────────────────────────────────────────────────────────

def generate_docx(data: dict, template: str = "modern") -> bytes:
    """Convert structured resume data to DOCX bytes using python-docx.

    Applies template-specific styling (font families, heading colors).
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Template-specific styling
    style_map = {
        "modern": {"font": "Calibri", "heading_color": RGBColor(0x37, 0x72, 0xFF), "accent": RGBColor(0x37, 0x72, 0xFF)},
        "classic": {"font": "Georgia", "heading_color": RGBColor(0x11, 0x11, 0x11), "accent": RGBColor(0x44, 0x44, 0x44)},
        "minimal": {"font": "Calibri", "heading_color": RGBColor(0x33, 0x33, 0x33), "accent": RGBColor(0x88, 0x88, 0x88)},
        "ats": {"font": "Times New Roman", "heading_color": RGBColor(0x00, 0x00, 0x00), "accent": RGBColor(0x00, 0x00, 0x00)},
    }
    style = style_map.get(template, style_map["modern"])

    # Set default font
    doc_style = doc.styles["Normal"]
    doc_style.font.name = style["font"]
    doc_style.font.size = Pt(10.5)

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    contact = data.get("contact", {})
    if isinstance(contact, dict):
        pass  # already a dict
    else:
        contact = {}

    # Name
    if contact.get("name"):
        p = doc.add_heading(contact["name"], level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.color.rgb = style["heading_color"]
            run.font.name = style["font"]

    # Contact info
    contact_items = []
    for key in ["email", "phone", "linkedin", "github", "portfolio"]:
        val = contact.get(key, "")
        if val:
            contact_items.append(str(val))
    if contact_items:
        p = doc.add_paragraph(" | ".join(contact_items))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style.font.size = Pt(9)

    def _add_section_heading(title: str):
        p = doc.add_heading(title, level=1)
        for run in p.runs:
            run.font.color.rgb = style["heading_color"]
            run.font.name = style["font"]
            run.font.size = Pt(12)

    # Summary
    if data.get("summary", "").strip():
        _add_section_heading("Summary")
        doc.add_paragraph(data["summary"])

    # Experience
    experience = data.get("experience", [])
    if experience:
        _add_section_heading("Experience")
        for exp in experience:
            title_line = exp.get("title", "")
            if exp.get("company"):
                title_line += f" — {exp['company']}"
            p = doc.add_paragraph()
            run = p.add_run(title_line)
            run.bold = True
            run.font.size = Pt(10.5)
            run.font.name = style["font"]

            date_line = ""
            if exp.get("start_date") or exp.get("end_date"):
                date_line = f'{exp.get("start_date", "")} – {exp.get("end_date", "")}'
            if exp.get("location"):
                date_line += f'  |  {exp["location"]}' if date_line else exp["location"]
            if date_line:
                p = doc.add_paragraph(date_line)
                p.style.font.size = Pt(9)

            for bullet in exp.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # Education
    education = data.get("education", [])
    if education:
        _add_section_heading("Education")
        for edu in education:
            degree_line = edu.get("degree", "")
            if edu.get("institution"):
                degree_line += f" — {edu['institution']}"
            p = doc.add_paragraph()
            run = p.add_run(degree_line)
            run.bold = True
            run.font.name = style["font"]

            extras = []
            if edu.get("year"):
                extras.append(edu["year"])
            if edu.get("location"):
                extras.append(edu["location"])
            if edu.get("gpa"):
                extras.append(f'GPA: {edu["gpa"]}')
            if extras:
                doc.add_paragraph(" | ".join(extras))

    # Skills
    skills = data.get("skills", [])
    if skills:
        _add_section_heading("Skills")
        for cat in skills:
            cat_name = cat.get("name", "")
            cat_skills = cat.get("skills", [])
            if cat_skills:
                p = doc.add_paragraph()
                run = p.add_run(f"{cat_name}: ")
                run.bold = True
                run.font.name = style["font"]
                p.add_run(", ".join(cat_skills)).font.name = style["font"]

    # Projects
    projects = data.get("projects", [])
    if projects:
        _add_section_heading("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.get("name", ""))
            run.bold = True
            run.font.name = style["font"]
            techs = proj.get("technologies", [])
            if techs:
                p.add_run(f'  ({", ".join(techs)})').font.name = style["font"]
            if proj.get("description"):
                doc.add_paragraph(proj["description"])
            for bullet in proj.get("bullets", []):
                doc.add_paragraph(bullet, style="List Bullet")

    # Certifications
    certs = data.get("certifications", [])
    if certs:
        _add_section_heading("Certifications")
        for c in certs:
            doc.add_paragraph(c, style="List Bullet")

    # Languages
    langs = data.get("languages", [])
    if langs:
        _add_section_heading("Languages")
        doc.add_paragraph(", ".join(langs))

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
